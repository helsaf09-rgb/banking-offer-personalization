from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCS = [
    ROOT / "docs" / "29_final_defense_presentation_en.md",
    ROOT / "docs" / "30_final_defense_talk_track_en.md",
    ROOT / "docs" / "31_final_defense_qna_en.md",
]
DEFAULT_OUTPUT_DIR = ROOT / "deliverables"


def set_run_font(run, size_pt: float = 12.0, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def clean_inline(text: str) -> str:
    return text.replace("`", "").replace("**", "").strip()


def add_markdown_to_doc(doc: Document, markdown: str) -> None:
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_inline(line[2:]))
            set_run_font(run, size_pt=16, bold=True)
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(clean_inline(line[3:]))
            set_run_font(run, size_pt=13, bold=True)
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(clean_inline(line[4:]))
            set_run_font(run, size_pt=12, bold=True)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            run = p.add_run(clean_inline(line[2:]))
            set_run_font(run)
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(clean_inline(line))
        set_run_font(run)


def build_docx(source: Path, output_dir: Path) -> Path:
    doc = Document()
    configure_doc(doc)
    markdown = source.read_text(encoding="utf-8")
    add_markdown_to_doc(doc, markdown)

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{source.stem}.docx"
    doc.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build final defense support DOCX files.")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("sources", type=Path, nargs="*", default=DEFAULT_DOCS)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    for source in args.sources:
        output_path = build_docx(source, args.output_dir)
        print(f"[ok] built {output_path}")


if __name__ == "__main__":
    main()
