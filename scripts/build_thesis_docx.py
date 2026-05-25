# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "25_thesis_final_ru.md"
DEFAULT_OUTPUT = ROOT / "deliverables" / "VKR_draft_gost.docx"
DEFAULT_STUDENT_NAME = "Сафонова Елена Михайловна"
DEFAULT_SUPERVISOR_NAME = "Ильвовский Дмитрий Алексеевич"
DEFAULT_DEGREE_PROGRAM = "01.04.02 Прикладная математика и информатика"
BODY_FONT_SIZE_PT = 14
TABLE_FONT_SIZE_PT = 12

DEFAULT_TITLE = (
    "Персонализация клиентских предложений на основе транзакционной "
    "активности пользователя банковских услуг"
)

SPECIAL_CENTER_HEADINGS = {
    "АННОТАЦИЯ",
    "КЛЮЧЕВЫЕ СЛОВА",
    "СПИСОК СОКРАЩЕНИЙ И УСЛОВНЫХ ОБОЗНАЧЕНИЙ",
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "ABSTRACT",
    "KEYWORDS",
    "ABBREVIATIONS",
    "LIST OF ABBREVIATIONS",
    "INTRODUCTION",
    "CONCLUSION",
    "REFERENCES",
}

FIGURE_TOKEN = re.compile(r"^\[\[FIGURE:\s*(.+?)\s*\]\]$")
PAGE_BREAK_TOKEN = re.compile(r"^\[\[PAGE_BREAK\]\]$")
TABLE_CAPTION_TOKEN = re.compile(r"^(Таблица|Table)\s+[A-Za-zА-Яа-я0-9.]+\s*[-–]\s*.+$")

FIGURE_IMAGE_MAP = {
    "Общая архитектура программного пайплайна персонализации клиентских предложений": "reports/thesis_figures/pipeline_architecture.png",
    "Структура синтетического датасета: пользователи, транзакции, офферы и взаимодействия": "reports/thesis_figures/synthetic_dataset_structure.png",
    "Схема построения пользовательского профиля и расчета итогового скоринга по офферам": "reports/thesis_figures/profile_scoring_scheme.png",
    "Скриншот Swagger-интерфейса API сервиса рекомендаций": "reports/thesis_figures/swagger_ui.png",
    "Пример ответа API с топом персонализированных банковских предложений": "reports/thesis_figures/api_response_example.png",
    "График сравнения NDCG@5 и MAP@5 для пяти реализованных моделей": "reports/thesis_figures/synthetic_model_metrics.png",
    "Сравнение результатов multi-seed benchmark для ключевых моделей": "reports/thesis_figures/multiseed_summary.png",
    "Скриншот структуры проекта в рабочем каталоге и основных модулей системы": "reports/thesis_figures/project_tree.png",
    "Структура проекта в рабочем каталоге и основные модули системы": "reports/thesis_figures/project_tree.png",
    "Скриншот ноутбука с разведочным анализом данных и основными графиками": "reports/thesis_figures/notebook_preview.png",
    "Фрагмент EDA-ноутбука с основными графиками": "reports/thesis_figures/notebook_preview.png",
    "Скриншот отчета с метриками baseline и advanced моделей": "reports/thesis_figures/model_report_preview.png",
    "Фрагмент отчета с метриками baseline и advanced моделей": "reports/thesis_figures/model_report_preview.png",
    "Диаграмма распределения количества транзакций по основным категориям": "reports/figures/eda_tx_count_by_category.png",
    "Диаграмма распределения пользователей по поведенческим сегментам": "reports/thesis_figures/segment_distribution.png",
    "График распределения суммарных расходов и медианных значений по сегментам": "reports/thesis_figures/segment_spend_summary.png",
    "Скриншот структуры папок проекта и основных директорий репозитория": "reports/thesis_figures/project_tree.png",
    "Структура папок проекта и основные директории репозитория": "reports/thesis_figures/project_tree.png",
    "Скриншот модуля модели и функций вычисления пользовательского профиля": "reports/thesis_figures/model_module_preview.png",
    "Фрагмент модуля модели и функций вычисления пользовательского профиля": "reports/thesis_figures/model_module_preview.png",
    "Скриншот PowerShell-скрипта полного прогона пайплайна": "reports/thesis_figures/run_all_script_preview.png",
    "Фрагмент PowerShell-скрипта полного прогона пайплайна": "reports/thesis_figures/run_all_script_preview.png",
    "Скриншот Swagger UI с описанием endpoint и параметров запроса": "reports/thesis_figures/swagger_ui.png",
    "Скриншот JSON-ответа сервиса с ранжированным списком рекомендаций": "reports/thesis_figures/api_response_example.png",
    "Фрагмент JSON-ответа сервиса с ранжированным списком рекомендаций": "reports/thesis_figures/api_response_example.png",
    "Скриншот отчета model comparison и итоговых таблиц метрик": "reports/thesis_figures/model_report_preview.png",
    "Фрагмент отчета model comparison и итоговых таблиц метрик": "reports/thesis_figures/model_report_preview.png",
    "Скриншот таблицы результатов подбора гиперпараметров time-decay модели": "reports/thesis_figures/time_decay_tuning.png",
    "Диаграмма результатов подбора гиперпараметров time-decay модели": "reports/thesis_figures/time_decay_tuning.png",
    "Скриншот bootstrap-анализа для сравнения time-decay и profile baseline": "reports/thesis_figures/bootstrap_ci_chart.png",
    "Диаграмма bootstrap-анализа для сравнения time-decay и profile baseline": "reports/thesis_figures/bootstrap_ci_chart.png",
    "Скриншот segment-wise метрик по основным поведенческим группам": "reports/thesis_figures/segment_metrics.png",
    "Тепловая карта segment-wise метрик по основным поведенческим группам": "reports/thesis_figures/segment_metrics.png",
    "Скриншот multi-seed summary report и усредненных значений NDCG@5": "reports/thesis_figures/multiseed_summary.png",
    "Диаграмма multi-seed summary report и усредненных значений NDCG@5": "reports/thesis_figures/multiseed_summary.png",
    "Скриншот запуска baseline pipeline и генерации синтетического датасета": "reports/thesis_figures/run_all_script_preview.png",
    "Фрагмент запуска baseline pipeline и генерации синтетического датасета": "reports/thesis_figures/run_all_script_preview.png",
    "Скриншот выполнения EDA-пайплайна и сформированного markdown-отчета": "reports/thesis_figures/notebook_preview.png",
    "Фрагмент выполнения EDA-пайплайна и сформированного markdown-отчета": "reports/thesis_figures/notebook_preview.png",
    "Скриншот таблицы сравнения моделей и итоговых метрик качества": "reports/thesis_figures/model_report_preview.png",
    "Фрагмент таблицы сравнения моделей и итоговых метрик качества": "reports/thesis_figures/model_report_preview.png",
    "Скриншот multi-seed benchmark с усредненными значениями метрик": "reports/thesis_figures/multiseed_summary.png",
    "Диаграмма multi-seed benchmark с усредненными значениями метрик": "reports/thesis_figures/multiseed_summary.png",
    "Скриншот окна Swagger UI после запуска локального FastAPI-сервиса": "reports/thesis_figures/swagger_ui.png",
    "Скриншот примера ответа endpoint recommend для конкретного пользователя": "reports/thesis_figures/api_response_example.png",
    "Фрагмент примера ответа endpoint recommend для конкретного пользователя": "reports/thesis_figures/api_response_example.png",
    "Overall architecture of the personalization pipeline": "reports/thesis_figures/pipeline_architecture.png",
    "Structure of the synthetic dataset: users, transactions, offers, and interactions": "reports/thesis_figures/synthetic_dataset_structure.png",
    "Distribution of transaction counts across the main categories": "reports/figures/eda_tx_count_by_category.png",
    "Distribution of users across behavioral segments": "reports/thesis_figures/segment_distribution.png",
    "Distribution of total spend and median values across segments": "reports/thesis_figures/segment_spend_summary.png",
    "User-profile construction and final offer scoring scheme": "reports/thesis_figures/profile_scoring_scheme.png",
    "Project structure and the main system modules": "reports/thesis_figures/project_tree.png",
    "Comparison of NDCG@5 and MAP@5 across the five implemented models": "reports/thesis_figures/synthetic_model_metrics.png",
    "Time-decay hyperparameter search results": "reports/thesis_figures/time_decay_tuning.png",
    "Segment-wise metrics across the main behavioral groups": "reports/thesis_figures/segment_metrics.png",
    "Bootstrap comparison of time-decay and the profile baseline": "reports/thesis_figures/bootstrap_ci_chart.png",
    "Comparison of multi-seed benchmark results for the key models": "reports/thesis_figures/multiseed_summary.png",
    "Swagger UI screenshot of the recommendation API": "reports/thesis_figures/swagger_ui.png",
    "Example API response with top personalized banking offers": "reports/thesis_figures/api_response_example.png",
}


def set_times_new_roman(target, size_pt: float | None = None, bold: bool | None = None) -> None:
    target.font.name = "Times New Roman"
    target._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    target.font.color.rgb = RGBColor(0, 0, 0)
    if size_pt is not None:
        target.font.size = Pt(size_pt)
    if bold is not None:
        target.font.bold = bold


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def set_a4_and_margins(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)


def apply_text_style(
    style,
    *,
    size_pt: float = BODY_FONT_SIZE_PT,
    bold: bool | None = None,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Cm(1.25),
    left_indent=None,
    line_spacing: float = 1.5,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    set_times_new_roman(style, size_pt, bold)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.first_line_indent = first_line_indent
    style.paragraph_format.alignment = alignment
    style.paragraph_format.space_before = Pt(space_before_pt)
    style.paragraph_format.space_after = Pt(space_after_pt)
    if left_indent is not None:
        style.paragraph_format.left_indent = left_indent


def configure_existing_style(doc: Document, style_name: str, **kwargs) -> None:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    apply_text_style(style, **kwargs)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    apply_text_style(normal)

    for style_name in ["List Number", "List Bullet"]:
        style = doc.styles[style_name]
        apply_text_style(
            style,
            size_pt=BODY_FONT_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Cm(0),
        )

    heading_1 = doc.styles["Heading 1"]
    apply_text_style(
        heading_1,
        size_pt=BODY_FONT_SIZE_PT,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Cm(0),
        space_before_pt=18,
        space_after_pt=12,
    )

    heading_2 = doc.styles["Heading 2"]
    apply_text_style(
        heading_2,
        size_pt=BODY_FONT_SIZE_PT,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=Cm(0),
        space_before_pt=12,
        space_after_pt=6,
    )

    if "Center Heading" not in [s.name for s in doc.styles]:
        center = doc.styles.add_style("Center Heading", WD_STYLE_TYPE.PARAGRAPH)
        apply_text_style(
            center,
            size_pt=BODY_FONT_SIZE_PT,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Cm(0),
        )

    if "Figure Caption" not in [s.name for s in doc.styles]:
        caption = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        apply_text_style(
            caption,
            size_pt=BODY_FONT_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Cm(0),
            space_before_pt=6,
            space_after_pt=6,
        )

    if "Table Caption" not in [s.name for s in doc.styles]:
        caption = doc.styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
        apply_text_style(
            caption,
            size_pt=BODY_FONT_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Cm(0),
            space_before_pt=6,
            space_after_pt=3,
        )

    if "Figure Placeholder" not in [s.name for s in doc.styles]:
        placeholder = doc.styles.add_style("Figure Placeholder", WD_STYLE_TYPE.PARAGRAPH)
        apply_text_style(
            placeholder,
            size_pt=BODY_FONT_SIZE_PT,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Cm(0),
        )
        placeholder.font.italic = True

    for style_name, left_indent in [("TOC 1", Cm(0)), ("TOC 2", Cm(0.75)), ("TOC 3", Cm(1.5)), ("toc 1", Cm(0)), ("toc 2", Cm(0.75)), ("toc 3", Cm(1.5))]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        apply_text_style(
            style,
            size_pt=BODY_FONT_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=Cm(0),
            left_indent=left_indent,
        )
        style.paragraph_format.tab_stops.add_tab_stop(
            Cm(16.2),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )

    configure_existing_style(
        doc,
        "диплом1",
        size_pt=BODY_FONT_SIZE_PT,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0),
        space_before_pt=12,
        space_after_pt=12,
    )
    configure_existing_style(
        doc,
        "доплом3",
        size_pt=BODY_FONT_SIZE_PT,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )


def add_page_number(section) -> None:
    section.different_first_page_header_footer = True
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE")


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_spacer(doc: Document, lines: int = 1, size_pt: float = 10.0) -> None:
    for _ in range(lines):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = paragraph.add_run(" ")
        set_times_new_roman(run, size_pt)


def add_title_page(
    doc: Document,
    title: str,
    language: str,
    student_name: str,
    supervisor_name: str,
    degree_program: str,
) -> None:
    def make_title_paragraph(alignment):
        paragraph = doc.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        return paragraph

    if language == "en":
        university = (
            "NATIONAL RESEARCH UNIVERSITY\n"
            "HIGHER SCHOOL OF ECONOMICS"
        )
        faculty = (
            "Faculty of Computer Science\n"
            'Master\'s Programme "Master of Data Science"'
        )
        work_type = "MASTER'S THESIS"
        topic_label = "on the topic"
        student_label = "Student"
        supervisor_label = "Supervisor"
        degree_label = "Degree programme"
        city_year = "Moscow\n2026"
    else:
        university = (
            "НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ\n"
            "«ВЫСШАЯ ШКОЛА ЭКОНОМИКИ»"
        )
        faculty = (
            "Факультет компьютерных наук\n"
            "Образовательная программа «Магистр наук о данных»"
        )
        work_type = "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА"
        topic_label = "на тему"
        student_label = "Студент"
        supervisor_label = "Научный руководитель"
        degree_label = "Направление подготовки"
        city_year = "Москва\n2026"

    p = make_title_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run(university).bold = True

    p = make_title_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run(faculty)

    add_spacer(doc, 3)

    p = make_title_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run(work_type).bold = True

    p = make_title_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run(topic_label)

    p = make_title_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run(f"«{title}»").bold = True

    add_spacer(doc, 4)

    p = make_title_paragraph(WD_ALIGN_PARAGRAPH.LEFT)
    p.add_run(f"{student_label}: {student_name}\n").bold = True
    p.add_run(f"{supervisor_label}: {supervisor_name}\n").bold = True
    p.add_run(f"{degree_label}: {degree_program}")

    add_spacer(doc, 3)

    p = make_title_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run(city_year)


def add_center_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Center Heading")
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def normalize_heading(text: str) -> str:
    return " ".join(text.upper().split())


def normalize_figure_caption(text: str) -> str:
    return " ".join(text.split())


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(TABLE_FONT_SIZE_PT)


def add_figure_placeholder(doc: Document, figure_label: str, caption: str, language: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Figure Placeholder"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    placeholder_text = (
        "Иллюстрация приведена в электронном приложении."
        if language == "ru"
        else "Figure is provided in the electronic appendix."
    )
    run = paragraph.add_run(placeholder_text)
    run.italic = True

    caption_paragraph = doc.add_paragraph(style="Figure Caption")
    caption_paragraph.add_run(f"{figure_label} - {caption}")


def add_figure_image(
    doc: Document,
    figure_label: str,
    caption: str,
    image_path: Path | None,
    language: str,
) -> None:
    if image_path is None or not image_path.exists():
        add_figure_placeholder(doc, figure_label, caption, language)
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))

    caption_paragraph = doc.add_paragraph(style="Figure Caption")
    caption_paragraph.add_run(f"{figure_label} - {caption}")


def add_plain_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Table Caption")
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)


def add_manual_numbered_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)


def add_heading_with_alignment(
    doc: Document,
    text: str,
    level: int,
    language: str,
    page_break_before: bool = False,
) -> None:
    heading = doc.add_heading(text, level=level)
    if level == 1 and normalize_heading(text) in SPECIAL_CENTER_HEADINGS:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.first_line_indent = Cm(0)
    heading.paragraph_format.page_break_before = page_break_before


def collect_figure_image_map(root: Path) -> dict[str, Path]:
    mapping: dict[str, Path] = {}
    for caption, relative_path in FIGURE_IMAGE_MAP.items():
        image_path = root / relative_path
        if image_path.exists():
            mapping[normalize_figure_caption(caption)] = image_path
    return mapping


def render_markdown_section(
    doc: Document,
    text: str,
    language: str,
    figure_prefix: str,
    figure_image_map: dict[str, Path],
    use_auto_figure_images: bool,
) -> None:
    lines = text.splitlines()
    i = 0
    first_section = True
    current_chapter: int | None = None
    figure_counter = 0

    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        if stripped.startswith("## "):
            page_break_before = not first_section
            first_section = False
            heading_text = clean_inline_markdown(stripped[3:])
            add_heading_with_alignment(
                doc,
                heading_text,
                level=1,
                language=language,
                page_break_before=page_break_before,
            )
            match = re.match(r"^(\d+)\.", heading_text)
            current_chapter = int(match.group(1)) if match else None
            figure_counter = 0
            i += 1
            continue

        if stripped.startswith("### "):
            heading_text = clean_inline_markdown(stripped[4:])
            add_heading_with_alignment(doc, heading_text, level=2, language=language)
            i += 1
            continue

        if PAGE_BREAK_TOKEN.match(stripped):
            doc.add_page_break()
            i += 1
            continue

        if TABLE_CAPTION_TOKEN.match(stripped):
            add_table_caption(doc, clean_inline_markdown(stripped))
            i += 1
            continue

        figure_match = FIGURE_TOKEN.match(stripped)
        if figure_match:
            figure_counter += 1
            if current_chapter is not None:
                label = f"{figure_prefix} {current_chapter}.{figure_counter}"
            else:
                label = f"{figure_prefix} {figure_counter}"

            caption = figure_match.group(1)
            normalized_caption = normalize_figure_caption(caption)
            if use_auto_figure_images and normalized_caption in figure_image_map:
                image_path = figure_image_map[normalized_caption]
                add_figure_image(doc, label, caption, image_path, language)
            else:
                add_figure_placeholder(doc, label, caption, language)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            separator = lines[i + 1].replace("|", "").strip()
            if separator and set(separator) <= {"-", ":", " "}:
                table_lines = [stripped]
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1

                rows = []
                for row_line in table_lines:
                    cells = [clean_inline_markdown(cell.strip()) for cell in row_line.strip("|").split("|")]
                    rows.append(cells)

                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for row_index, row in enumerate(rows):
                    for col_index, value in enumerate(row):
                        table.cell(row_index, col_index).text = value
                style_table(table)
                add_spacer(doc, lines=1, size_pt=4.0)
                continue

        if re.match(r"^\d+\.\s", stripped):
            add_manual_numbered_paragraph(doc, clean_inline_markdown(stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline_markdown(stripped[2:]))
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith(("#", "-", "|")) or re.match(r"^\d+\.\s", nxt) or FIGURE_TOKEN.match(nxt) or PAGE_BREAK_TOKEN.match(nxt):
                break
            paragraph_lines.append(nxt)
            i += 1

        add_plain_paragraph(doc, clean_inline_markdown(" ".join(paragraph_lines)))


def build_document(
    source_path: Path,
    output_path: Path,
    title: str,
    language: str,
    figure_prefix: str,
    use_auto_figure_images: bool,
    student_name: str,
    supervisor_name: str,
    degree_program: str,
) -> Path:
    if not source_path.exists():
        raise FileNotFoundError(f"Source markdown does not exist: {source_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_a4_and_margins(doc.sections[0])
    configure_styles(doc)
    enable_update_fields_on_open(doc)
    add_page_number(doc.sections[0])
    add_title_page(
        doc,
        title=title,
        language=language,
        student_name=student_name,
        supervisor_name=supervisor_name,
        degree_program=degree_program,
    )

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_a4_and_margins(main_section)
    add_page_number(main_section)

    toc_title = "СОДЕРЖАНИЕ" if language == "ru" else "CONTENTS"
    add_center_heading(doc, toc_title)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')

    figure_image_map = collect_figure_image_map(ROOT) if use_auto_figure_images else {}
    text = source_path.read_text(encoding="utf-8-sig")
    render_markdown_section(
        doc=doc,
        text=text,
        language=language,
        figure_prefix=figure_prefix,
        figure_image_map=figure_image_map,
        use_auto_figure_images=use_auto_figure_images,
    )

    doc.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build thesis .docx from markdown source.")
    parser.add_argument("--source", default=str(DEFAULT_SOURCE), help="Path to markdown source.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="Path to output .docx.")
    parser.add_argument("--title", default=DEFAULT_TITLE, help="Title on the title page.")
    parser.add_argument("--language", choices=["ru", "en"], default="ru", help="Document language.")
    parser.add_argument("--figure-prefix", default="\u0420\u0438\u0441\u0443\u043d\u043e\u043a", help="Prefix for figure captions.")
    parser.add_argument(
        "--student-name",
        default=DEFAULT_STUDENT_NAME,
        help="Student full name for the title page.",
    )
    parser.add_argument(
        "--supervisor-name",
        default=DEFAULT_SUPERVISOR_NAME,
        help="Supervisor full name for the title page.",
    )
    parser.add_argument(
        "--degree-program",
        default=DEFAULT_DEGREE_PROGRAM,
        help="Degree program for the title page.",
    )
    parser.add_argument(
        "--auto-figure-images",
        action="store_true",
        help="Auto-insert png figures from reports folders into [[FIGURE: ...]] blocks.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    output = build_document(
        source_path=Path(args.source),
        output_path=Path(args.output),
        title=args.title,
        language=args.language,
        figure_prefix=args.figure_prefix,
        use_auto_figure_images=args.auto_figure_images,
        student_name=args.student_name,
        supervisor_name=args.supervisor_name,
        degree_program=args.degree_program,
    )
    print(f"[ok] thesis docx saved: {output}")


